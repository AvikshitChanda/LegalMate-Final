from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from typing import Literal
from dataclasses import dataclass
from langchain_groq import ChatGroq
from llamaapi import LlamaAPI
from langchain_experimental.llms import ChatLlamaAPI
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain.prompts import PromptTemplate
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ChatMessageHistory, ConversationBufferMemory
from docx import Document
import os
import json
import random

# Initialize Flask app and CORS
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:5173"}})

@dataclass
class Message:
    """Class for keeping track of a chat message."""
    origin: Literal["👤 Human", "👨🏻‍⚖️ Ai"]
    message: str

# Download Hugging Face embeddings
def download_hugging_face_embeddings():
    embeddings = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')
    return embeddings

# Initialize the LLM-based conversation system
def initialize_conversation():
    llama = LlamaAPI("d9f5e52518c6432f9c568a308d17554b767f2543a040477d9888bf9704d4d065")
    model = ChatLlamaAPI(client=llama)
    chat = ChatGroq(
        temperature=0.5,
        groq_api_key="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",
        model_name="llama3-70b-8192"
    )

    embeddings = download_hugging_face_embeddings()

    persist_directory = "./chroma-db"
    chroma_client = Chroma(
        collection_name="legal-advisor",
        embedding_function=embeddings,
        persist_directory=persist_directory
    )

    # Define prompt template for legal context
    prompt_template = """
        You are an Indian legal expert helping users generate marital settlement agreements. Based on the data in the vector database and the values given by the user and their religion, decide which marriage act to consider and fill all the values for:
        - Amount for Basic Payment Provisions
        - Amount under Basic Payment Provision (Spousal Support)
        - Termination of Jurisdiction (Husband/Wife)
        - Amount Payment to Balance Division
        - Amount for Maintenance Repairs
        - Percentage for Husband and Wife Tax Refunds under Allocation of Income Tax
        - Assets allocated to husband
        - Assets allocated to wife
        - Marriage Act
        Return only the values in JSON format.
        Context: {context}
        Question: {question}
        Answer:
    """

    PROMPT = PromptTemplate(template=prompt_template, input_variables=["context", "question"])

    # Set up message history and conversation memory
    message_history = ChatMessageHistory()
    memory = ConversationBufferMemory(
        memory_key="chat_history",
        output_key="answer",
        chat_memory=message_history,
        return_messages=True,
    )

    # Create a conversational retrieval chain
    retrieval_chain = ConversationalRetrievalChain.from_llm(
        llm=chat,
        chain_type="stuff",
        retriever=chroma_client.as_retriever(search_kwargs={'k': 2}),
        return_source_documents=True,
        combine_docs_chain_kwargs={"prompt": PROMPT},
        memory=memory
    )

    return retrieval_chain
def generate_marital_settlement_agreement(values, user_inputs):
    doc = Document()
    
    doc.add_heading('Marital Settlement Agreement', 0)

    # Introduction section
    doc.add_paragraph(f"This Marital Settlement Agreement is made and entered , between the following parties: ")
    doc.add_paragraph(
    f"Mr. {user_inputs['Husband_Name']}, hereinafter referred to as the 'Husband', is a person of legal capacity and sound mind. This document serves as a formal acknowledgment of the intentions agreed upon by both parties, focusing on establishing a clear framework for their financial responsibilities and marital arrangements. The Husband and the Wife understand that this agreement is essential for promoting transparency and protecting their respective interests, allowing both parties to move forward with confidence and clarity."
)
    doc.add_paragraph(
    f"Mrs. {user_inputs['Wife_Name']}, hereinafter referred to as the 'Wife', is an individual of sound mind and legal capacity. This agreement signifies the mutual understanding between the Wife and Mr. {user_inputs['Husband_Name']}, detailing how their marital and financial affairs will be resolved. Both parties recognize the importance of this agreement in fostering clarity and ensuring that their rights and obligations are addressed fairly. It is their intention to create a harmonious resolution that reflects their shared values and mutual respect."
)

    # Recitals
    doc.add_heading('Recitals', level=1)
    doc.add_paragraph(f"WHEREAS, the parties were lawfully married on {user_inputs['Marriage_Date']} under the {values.get('Marriage Act', 'N/A')}, a lawfully recognized and binding union, signifying their commitment to each other both legally and socially.")
    doc.add_paragraph(f"WHEREAS, the Husband and Wife have acknowledged their marital challenges, and after considerable reflection, they have been living apart since {user_inputs['Separation_Date']} due to irreconcilable differences. This period of separation is recognized by both parties as a prelude to this final settlement.")
    doc.add_paragraph(f"WHEREAS, the parties mutually desire to resolve all financial, legal, and personal obligations amicably, without resorting to litigation, and thus agree to the terms of this document in good faith.")

    # Financial Arrangements
    doc.add_heading('Financial Arrangements', level=1)
    doc.add_paragraph(f"During the course of the marriage, both parties have contributed to the household financially. The Husband has earned an annual income of ₹{user_inputs['Husband_Income']}, which has been derived from his professional employment, as well as additional sources such as investments or businesses.")
    doc.add_paragraph(f"The Wife, equally contributing to the marital partnership, has earned an annual income of ₹{user_inputs['Wife_Income']}, which includes her primary employment and any freelance work or other business ventures she may have undertaken during the marriage.")
    doc.add_paragraph(f"Both parties recognize the shared responsibilities they upheld throughout the marriage, including the management of household expenses, investments, and financial planning. This agreement seeks to ensure that both parties retain financial independence while fulfilling any obligations laid out herein.")

    # Basic Payment Provisions
    doc.add_heading('Basic Payment Provisions', level=1)
    doc.add_paragraph(f"The Husband agrees to pay a sum of ₹{values.get('Amount for Basic Payment Provisions', 'N/A')} to the Wife. This sum will cover any basic expenses, ensuring that the Wife maintains a standard of living comparable to the one she had during the marriage.")
    doc.add_paragraph(f"This payment will be made on a regular basis as outlined in a separate payment schedule, which both parties agree upon. The aim of this provision is to provide financial support to the Wife while she adjusts to post-marital life.")
    doc.add_paragraph(f"Both parties agree that this payment is fair and equitable, considering the financial circumstances and contributions made by both the Husband and Wife during the marriage, as well as the future financial needs of the Wife.")

    # Spousal Support
    doc.add_heading('Spousal Support', level=1)
    doc.add_paragraph(f"In addition to the basic payment provisions, the Husband will provide spousal support in the form of ₹{values.get('Amount under Basic Payment Provision (Spousal Support)', 'N/A')}. This amount has been calculated based on the Wife’s financial needs, her earning potential, and the standard of living enjoyed during the marriage.")
    doc.add_paragraph(f"The spousal support is aimed at helping the Wife maintain financial stability, while she navigates her post-marital circumstances and transitions into financial independence. Both parties have mutually agreed on this provision.")
    doc.add_paragraph(f"The parties acknowledge that this spousal support is intended to continue for a predefined period or until certain conditions are met, such as the remarriage of the Wife or significant changes in her financial situation.")

    # Termination of Jurisdiction
    doc.add_heading('Termination of Jurisdiction', level=1)
    doc.add_paragraph(f"Both parties agree that jurisdiction over spousal support, including its modification or termination, shall be governed by this Agreement. As such, both Husband and Wife acknowledge that upon the fulfillment of the spousal support obligations, the court’s jurisdiction over spousal support will cease.")
    doc.add_paragraph(f"The Husband’s obligation to provide spousal support will terminate upon either the expiration of the agreed period or the occurrence of specific events such as the Wife’s remarriage or a significant increase in her income.")
    doc.add_paragraph(f"Both parties agree that these terms are fair and reasonable given their financial circumstances and the desire to avoid future disputes or modifications regarding spousal support.")

    # Payment for Balance Division
    doc.add_heading('Payment to Balance Division', level=1)
    doc.add_paragraph(f"To ensure an equitable division of marital assets, the Husband will pay ₹{values.get('Amount Payment to Balance Division', 'N/A')} to the Wife. This payment reflects the agreed division of assets and is aimed at providing a fair financial balance between both parties.")
    doc.add_paragraph(f"This division considers both the tangible and intangible contributions made by each party throughout the course of the marriage. The amount has been carefully calculated to ensure both parties leave the marriage on equal financial footing.")
    doc.add_paragraph(f"The division of assets, including the payment, will be completed within the timeline agreed upon by both parties and set forth in this Agreement. Any failure to make timely payments will result in penalties as described under the applicable sections of the law.")

    # Maintenance and Repairs
    doc.add_heading('Maintenance and Repairs', level=1)
    doc.add_paragraph(f"The Husband will also contribute to the maintenance and repairs of any joint property or assets that are in the process of being divided. The agreed contribution is ₹{values.get('Amount for Maintenance Repairs', 'N/A')}, which will be used for necessary repairs or upkeep until the asset division is finalized.")
    doc.add_paragraph(f"Both parties acknowledge the need to maintain the value of their joint assets and agree that these funds will be used specifically for maintenance purposes. The goal is to ensure that neither party suffers financial loss due to neglect of the shared assets.")
    doc.add_paragraph(f"The exact allocation of these funds and the specific repairs or maintenance tasks will be decided jointly by the parties to prevent any future disputes regarding the proper use of the funds.")

    # Tax Refund Allocation
    doc.add_heading('Allocation of Income Tax', level=1)
    doc.add_paragraph(f"Both Husband and Wife agree to allocate the percentage of tax refunds in the following manner: {values.get('Percentage for Husband and Wife Tax Refunds under Allocation of Income Tax', 'N/A')} for the Husband and the remainder for the Wife.")
    doc.add_paragraph(f"This allocation reflects the income and tax responsibilities borne by both parties during the course of the marriage, and ensures that the distribution of refunds is equitable and fair.")
    doc.add_paragraph(f"The parties agree to file their taxes jointly or separately as per the circumstances, and ensure that all tax-related responsibilities are fulfilled in accordance with applicable laws.")

    # Asset Division
    doc.add_heading('Asset Division', level=1)
    doc.add_paragraph(f"The total assets accumulated during the marriage amount to ₹{user_inputs['total_assets']}. These assets shall be divided equitably between both parties as outlined in this Agreement.")
    doc.add_paragraph(f"The Husband will receive the following percentage of the assets: {values.get('Assets allocated to husband', 'N/A')}, while the Wife will receive the remaining percentage as outlined: {values.get('Assets allocated to wife', 'N/A')}.")
    doc.add_paragraph(f"Both parties acknowledge that this division is fair and have agreed to it willingly, without coercion or undue influence. The division shall be completed by 20/12/22 and all related legal documentation shall be finalized accordingly.")

    # Save the document to a file
    random_number = random.randint(1000, 9999)
    filename = f'Marital_Settlement_Agreement_{random_number}.docx'
    doc.save(filename)
    
    return filename


@app.route('/generate-agreement', methods=['POST'])
def generate_agreement():
    # Receive user input from frontend
    user_inputs = request.json

    # Initialize conversation (LLM-based)
    conversation = initialize_conversation()

    # Prepare the financial context
    financial_context = f"""
    Husband's Income: {user_inputs['Husband_Income']}
    Wife's Income: {user_inputs['Wife_Income']}
    Husband's Religion: {user_inputs['Husband_Religion']}
    Wife's Religion: {user_inputs['Wife_Religion']}
    Total Assets: {user_inputs['total_assets']}
    """

    # Formulate the prompt for the LLM
    prompt = f"""
    Context: {financial_context}
    Fill in the following values for the marital settlement agreement:
    - Amount for Basic Payment Provisions
    - Amount under Basic Payment Provision (Spousal Support)
    - Termination of Jurisdiction (Husband/Wife)
    - Amount Payment to Balance Division
    - Amount for Maintenance Repairs
    - Percentage for Husband and Wife Tax Refunds under Allocation of Income Tax
    - Assets allocated to husband
    - Assets allocated to wife
    - Marriage Act
    """

    # Call the LLM and get response
    response = conversation({"question": prompt})
    generated_values_str = response['answer']

    # Extract JSON from the response
    st = generated_values_str.find('{')
    end = generated_values_str.rfind('}')
    if st == -1 or end == -1:
        return jsonify({"error": "Invalid JSON response from model"}), 400

    json_str = generated_values_str[st:end + 1]
    try:
        values = json.loads(json_str)
    except json.JSONDecodeError as e:
        return jsonify({"error": f"JSON decoding error: {e}"}), 400

    # Generate the document based on provided values
    filename = generate_marital_settlement_agreement(values, user_inputs)

    # Send the generated document as a file download
    return send_file(filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
