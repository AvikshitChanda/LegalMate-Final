from typing import Literal
from dataclasses import dataclass
from langchain_groq import ChatGroq
from llamaapi import LlamaAPI
from langchain_experimental.llms import ChatLlamaAPI
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain.prompts import PromptTemplate
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ChatMessageHistory, ConversationBufferMemory
from docx import Document
import os
import json
import re
import random

@dataclass
class Message:
    """Class for keeping track of a chat message."""
    origin: Literal["👤 Human", "👨🏻‍⚖️ Ai"]
    message: str

# Download and load Hugging Face embeddings
def download_hugging_face_embeddings():
    embeddings = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')
    return embeddings

# Initialize the conversation system with LlamaAPI and Groq
def initialize_conversation():
    llama = LlamaAPI("d9f5e52518c6432f9c568a308d17554b767f2543a040477d9888bf9704d4d065")
    model = ChatLlamaAPI(client=llama)
    chat = ChatGroq(
        temperature=0.5,
        groq_api_key="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",
        model_name="llama3-70b-8192"
    )

    embeddings = download_hugging_face_embeddings()

    persist_directory = "./chroma-db"
    chroma_client = Chroma(
        collection_name="legal-advisor",
        embedding_function=embeddings,
        persist_directory=persist_directory
    )

    prompt_template = """
        You are an Indian legal expert helping users generate marital settlement agreements. Based on the data in the vector database and the values given by the user and their religion, decide which marriage act to consider and fill all the values for:
        - Amount for Basic Payment Provisions
        - Amount under Basic Payment Provision (Spousal Support)
        - Termination of Jurisdiction (Husband/Wife)
        - Amount Payment to Balance Division
        - Amount for Maintenance Repairs
        - Percentage for Husband and Wife Tax Refunds under Allocation of Income Tax
        - Splitting of assets from total assets
        - Mention the marriage act and the reason behind it
        
        Return only the values in JSON format.
        Context: {context}
        Question: {question}
        Answer:
    """

    PROMPT = PromptTemplate(template=prompt_template, input_variables=["context", "question"])

    message_history = ChatMessageHistory()
    memory = ConversationBufferMemory(
        memory_key="chat_history",
        output_key="answer",
        chat_memory=message_history,
        return_messages=True,
    )

    retrieval_chain = ConversationalRetrievalChain.from_llm(
        llm=chat,
        chain_type="stuff",
        retriever=chroma_client.as_retriever(search_kwargs={'k': 2}),
        return_source_documents=True,
        combine_docs_chain_kwargs={"prompt": PROMPT},
        memory=memory
    )

    return retrieval_chain

def generate_marital_settlement_agreement(values, user_inputs):
    doc = Document()
        
    doc.add_heading('Marital Settlement Agreement', 0)

    doc.add_paragraph(f"This Marital Settlement Agreement (the 'Agreement') is made on [Date] between:")
    doc.add_paragraph(f"{user_inputs['Husband_Name']}, son of [Husband's Father's Name], residing at [Husband's Address], hereinafter referred to as 'Husband';")
    doc.add_paragraph(f"{user_inputs['Wife_Name']}, daughter of [Wife's Father's Name], residing at [Wife's Address], hereinafter referred to as 'Wife'.")

    doc.add_heading('Recitals', level=1)
    doc.add_paragraph(f"The Parties were married on {user_inputs['Marriage_Date']} under the provisions of the {values.get('Marriage Act', 'N/A')}. This was a legally binding union, which implied various legal, emotional, and social responsibilities shared between both parties. They agreed to share their lives, financial burdens, and emotional support through the bond of marriage, acknowledging that marriage is a partnership involving both rights and duties.")

    doc.add_paragraph(f"However, despite their mutual efforts, the marriage faced challenges. The Parties have been living separately since {user_inputs['Separation_Date']}, having experienced irreconcilable differences that led to their decision to separate. This Agreement reflects the Parties' desire to resolve their financial, legal, and personal matters amicably and with respect for one another, avoiding litigation or any adversarial process that may cause further stress or emotional strain.")

    doc.add_heading('Financial Arrangements', level=1)
    doc.add_paragraph(f"The financial obligations and contributions during the marriage have been significant. The Husband's annual income is ₹{user_inputs['Husband_Income']}. This income is derived from multiple sources, including his primary employment, business ventures, investments, and other financial undertakings. It represents his financial capacity to support both himself and the Wife during the marriage. The Husband has played a key role in managing certain household expenses, investments, and long-term financial planning.")

    doc.add_paragraph(f"The Wife's annual income is ₹{user_inputs['Wife_Income']}, which includes her salary and any other financial sources, such as freelance work, investments, or other business-related income. The Wife has been equally committed to the financial stability of the marriage, contributing towards shared household expenses, savings, and any dependents they may have. This Agreement seeks to ensure that both parties continue to enjoy financial independence while meeting their obligations as determined.")

    doc.add_heading('Basic Payment Provisions', level=1)
    doc.add_paragraph(f"The Husband agrees to pay a total sum of ₹{values.get('Amount for Basic Payment Provisions', 'N/A')} as part of the marital settlement. This payment will cover essential support that the Wife may require post-separation, ensuring that her living standards are maintained at a reasonable level, considering their shared financial history. The payment is intended to reflect not only financial compensation but also the acknowledgment of the shared responsibilities both parties undertook during their marriage, such as mortgage payments, household maintenance, and savings for future needs.")

    doc.add_paragraph(f"The payment will be made in a manner and schedule agreed upon by both parties. The exact dates, amounts, and form of payment (e.g., bank transfer, cheque, etc.) will be clearly outlined to avoid any disputes. This provision ensures that both parties can uphold their financial commitments while transitioning to their post-marital lives.")

    doc.add_heading('Spousal Support', level=1)
    doc.add_paragraph(f"The Husband shall provide spousal support to the Wife amounting to ₹{values.get('Amount under Basic Payment Provision (Spousal Support)', 'N/A')} per month. This spousal support is a recognition of the sacrifices and contributions made by the Wife during the marriage, including, but not limited to, managing the household, contributing to shared assets, and possibly foregoing certain career opportunities for the benefit of the family. The support payments are designed to ensure that the Wife can maintain a standard of living similar to what was enjoyed during the marriage, or at least sufficient for her needs.")

    doc.add_paragraph(f"Spousal support will continue for a predetermined period, allowing the Wife to transition smoothly to financial independence. The duration and amount of this support can be adjusted, depending on any significant changes in the financial or personal circumstances of either party, such as changes in employment, health, or the financial market conditions. Any such modifications will be subject to mutual agreement or judicial intervention, if necessary.")

    doc.add_heading('Division of Assets', level=1)
    doc.add_paragraph(f"The division of assets is a crucial aspect of the marital settlement. The Parties agree to divide their assets as follows, ensuring that both receive a fair share of the wealth accumulated during their marriage. This includes not only tangible assets but also intangible assets, such as investments, retirement funds, and savings.")

    # Expanding asset lists
    husband_assets = ', '.join([f"{asset['name']} (₹{asset['value']})" for asset in values.get('Assets allocated to Husband', [])])
    wife_assets = ', '.join([f"{asset['name']} (₹{asset['value']})" for asset in values.get('Assets allocated to Wife', [])])

    doc.add_paragraph(f"Assets allocated to Husband: {husband_assets if husband_assets else 'None'}. These may include assets such as real estate, business interests, vehicles, personal belongings, and other valuable items that the Husband has an interest in. The allocation of these assets is determined based on several factors, including ownership, contribution towards the purchase, and the respective needs of each party.")
    doc.add_paragraph(f"Assets allocated to Wife: {wife_assets if wife_assets else 'None'}. These assets may include savings accounts, personal investments, jewelry, or other valuable personal property. The division takes into account both the financial and emotional value attached to these assets, ensuring that the Wife receives a portion that reflects her contributions and interests.")

    doc.add_heading('Tax Refunds', level=1)
    doc.add_paragraph(f"The percentage of tax refunds allocated between the Parties will be {values.get('Percentage for Husband and Wife Tax Refunds under Allocation of Income Tax', 'N/A')}% for the Husband and {values.get('Tax_Refund_Percentage_Wife', 'N/A')}% for the Wife. This allocation is designed to reflect the proportional contribution of each party to the household's overall income and tax liability during the marriage.")

    doc.add_paragraph(f"Tax refunds from previous years, as well as any pending refunds, will be divided accordingly. The tax filing process for the duration of the marriage has been handled jointly, and this Agreement ensures that both parties receive their fair share of any tax benefits, including refunds, deductions, and credits. This also includes the handling of any future tax liabilities or audits, which will be addressed based on the financial agreements laid out in this Agreement.")

    # Save the document to a file
    # Generate a unique filename by adding a random number
    random_number = random.randint(1000, 9999)
    filename = f'Marital_Settlement_Agreement_{random_number}.docx'
    doc.save(filename)
    return filename


# Function to generate the document based on LLM responses
def generate_document(conversation, user_inputs):
    # Prepare the financial context for LLM
    financial_context = f"""
    Husband's Income: {user_inputs['Husband_Income']}
    Wife's Income: {user_inputs['Wife_Income']}
    Husband's Religion: {user_inputs['Husband_Religion']}
    Wife's Religion: {user_inputs['Wife_Religion']}
    Total Assets: {user_inputs['total_assets']}
    """

    # Combine the financial context and question into one prompt
    prompt = f"""
    Context: {financial_context}
    Fill in the following values for the marital settlement agreement:
    - Amount for Basic Payment Provisions
    - Amount under Basic Payment Provision (Spousal Support)
    - Termination of Jurisdiction (Husband/Wife)
    - Amount Payment to Balance Division
    - Amount for Maintenance Repairs
    - Percentage for Husband and Wife Tax Refunds under Allocation of Income Tax
    - Assets allocated to husband
    - Assets allocated to wife
    - Marriage Act
    """

    # Call the conversation chain and get the answer
    response = conversation({"question": prompt})
    generated_values_str = response['answer']  # Extract only the answer

    # Extract JSON from the response
    s = generated_values_str
    st = s.find('{')  # Find the starting index of JSON
    end = s.rfind('}')  # Find the ending index of JSON
    if st == -1 or end == -1:
        print("No valid JSON found in the response.")
        return

    json_str = s[st:end + 1]  # Extract the JSON string
    #print(json_str)
    # Parse the extracted JSON string into a dictionary
    try:
        values = json.loads(json_str)  # Parse the JSON response
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON: {e}")
        return  # Exit if the response is not valid JSON

    # Generate the document with the provided values
    print(values['Amount for Basic Payment Provisions'])
    filename = generate_marital_settlement_agreement(values, user_inputs)
    print(f"Document generated and saved as: {filename}")

# Example default inputs for testing
def get_default_user_details():
    return {
        "Husband_Name": "John Doe",
        "Wife_Name": "Jane Doe",
        "Husband_Religion": "Christian",
        "Wife_Religion": "Christian",
        "Marriage_Date": "2010-05-15",
        "Separation_Date": "2022-08-01",
        "Number_of_Children": 2,
        "Children": [
            {"name": "Alice Doe", "dob": "2012-03-20"},
            {"name": "Bob Doe", "dob": "2015-06-10"}
        ],
        "Husband_Income": 387332.00,
        "Wife_Income": 187040.00,
        "total_assets": [
            {"name": "House", "value": 500000.00},
            {"name": "Car", "value": 30000.00},
            {"name": "Stocks", "value": 150000.00}
        ]
    }

if __name__ == "__main__":
    conversation = initialize_conversation()
    user_inputs = get_default_user_details()
    generate_document(conversation, user_inputs)
